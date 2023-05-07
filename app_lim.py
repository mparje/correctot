import io
import openai
import streamlit as st
import docx

# Configure OpenAI API key
api_key = st.sidebar.text_input("Enter your OpenAI API key", type="password")

st.sidebar.title("App Description")
st.sidebar.write(
    """
    This Streamlit application demonstrates the use of GPT-3.5-turbo, an 
    advanced AI language model, to automatically correct grammar
    and improve style in Word documents. To use the application, upload a .docx 
    file and let GPT-3.5-turbo process each paragraph of the document.
    Once processed, download the corrected document in .docx format. NOTE: this is a limited version.
    """
)

custom_prompt = st.sidebar.text_input("Enter your custom prompt (Leave empty for default prompt)")

if not api_key:
    st.warning("Please enter a valid API key to continue.")
else:
    openai.api_key = api_key

def gpt_correct_prompt(prompt):
    completions = openai.Completion.create(engine="text-davinci-003", prompt=prompt, max_tokens=1024, n=1, stop=None,
                                           temperature=0.5)
    message = completions.choices[0].text.strip()
    return message

def process_document(doc_buffer, custom_prompt, max_paragraphs=10):
    doc = docx.Document(doc_buffer)
    corrected_doc = docx.Document()

    paragraph_count = 0

    for paragraph in doc.paragraphs:
        if paragraph_count >= max_paragraphs:
            break

        original_text = paragraph.text
        if original_text.strip() == "":
            corrected_doc.add_paragraph()
            continue

        if custom_prompt:
            prompt = f"{custom_prompt}\n\nOriginal text:\n'{original_text}'\n\nCorrected text:"
        else:
            prompt = f"Rewrite the following paragraph, correcting grammatical errors and improving the style:\n'{original_text}'\n\nCorrected text:"
            
        corrected_text = gpt_correct_prompt(prompt)
        corrected_paragraph = corrected_doc.add_paragraph(corrected_text)
        for run in paragraph.runs:
            if run.bold:
                corrected_paragraph.runs[-1].bold = True
            if run.italic:
                corrected_paragraph.runs[-1].italic = True
            if run.underline:
                corrected_paragraph.runs[-1].underline = True

        paragraph_count += 1

    return corrected_doc

st.image('http://tareas.site/wp-content/uploads/2023/05/logo-e1683457433611.png')
st.write('Upload a Word document (.docx) and let GPT-3.5-turbo correct up to 10 paragraphs with grammatical errors and style.')

uploaded_file = st.file_uploader("Upload file", type=[".docx"])

if uploaded_file is not None:
    with io.BytesIO(uploaded_file.getvalue()) as doc_buffer:
        corrected_doc = process_document(doc_buffer, custom_prompt)

        with io.BytesIO() as bytes_io:
            corrected_doc.save(bytes_io)
            bytes_io.seek(0)
            st.download_button(label="Download corrected document", data=bytes_io,
                               file_name="corrected_document.docx")
